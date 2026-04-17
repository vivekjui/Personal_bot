import json
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_noting_text(text):
    if not text:
        return ""
    
    # Remove Note# prefix
    text = re.sub(r'(?i)Note\s*#\s*\d+[.\d]*\s*', '', text)
    
    # Remove signature placeholders and dates
    patterns = [
        r'\(हस्ताक्षर\)', r'\(नाम\)', r'\(पदनाम\)', r'\(विभाग\)',
        r'\(Signature\)', r'\(Name\)', r'\(Designation\)', r'\(Department\)',
        r'दिनांक\s*:?.*', r'Date\s*:?.*',
        r'\d{2}/\d{2}/\d{4}\s+\d{2}:\d{2}\s+[AP]M', # Timestamp 07/04/2025 03:28 PM
        r'VIVEK KUMAR',
        r'Store Keeper'
    ]
    for p in patterns:
        text = re.sub(p, '', text, flags=re.IGNORECASE)
    
    # Clean up trailing whitespace and multiple newlines
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()

def export_to_docx():
    # Source file
    base_file = Path("d:/APMD_eOffice_Bot/Vivek_Kumar_Notings_Categorized.jsonl")
    if not base_file.exists():
        print(f"Error: {base_file} not found")
        return

    doc = Document()
    doc.add_heading('Vivek Kumar Noting Library - Cleaned for RAG', 0)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Stage (LLM Label)'
    hdr_cells[1].text = 'Sanitized Content'
    
    # Set bold headers
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    count = 0
    with open(base_file, "r", encoding="utf-8") as f:
        for line in f:
            if not line.strip(): continue
            try:
                item = json.loads(line)
                stage = item.get("category", "General")
                original_text = item.get("text", "")
                cleaned_text = clean_noting_text(original_text)
                
                if cleaned_text:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(stage)
                    row_cells[1].text = cleaned_text
                    count += 1
            except Exception as e:
                print(f"Error processing line: {e}")

    output_path = "d:/APMD_eOffice_Bot/Vivek_Kumar_Notings_Clean_Table.docx"
    doc.save(output_path)
    print(f"Successfully exported {count} notings to {output_path}")

if __name__ == "__main__":
    export_to_docx()
