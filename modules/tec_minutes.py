import os
import json
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Path to the learned patterns
LEARNED_PATTERNS_PATH = os.path.join(os.path.dirname(__file__), "..", "tec_patterns_learned.json")

# Default Committee Members 2025-26
DEFAULT_COMMITTEE = {
    "chairman": "Shri. Lalit Mohan Singh Maura, DDG & HoD",
    "member_finance": "Shri. Pavan Kumar, Director (G)",
    "member_secretary": "Shri. Pakki Varaprasad (SE)",
    "member_indenting": "....." # Placeholder for variable member
}

def load_learned_patterns():
    if os.path.exists(LEARNED_PATTERNS_PATH):
        with open(LEARNED_PATTERNS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def extract_entities_from_raw_text(text):
    """
    Uses regex and simple patterns to extract tender details.
    Will be supplemented by LLM in the main drafting call.
    """
    patterns = {
        "file_number": r"(E-\d+:\s*[A-Z0-9/-]+)",
        "gem_bid_id": r"(GEM/\d+/[A-Z]/\d+)",
        "bid_date": r"(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
    }
    
    entities = {}
    for key, pattern in patterns.items():
        match = re.search(pattern, text)
        if match:
            entities[key] = match.group(1)
            
    return entities

def create_tec_docx(output_path, content_html, title="TEC Minutes"):
    """
    Generates a DOCX file with LEGAL size and 1-inch margins.
    Uses simple HTML to DOCX conversion for basic formatting.
    """
    doc = Document()
    
    # Set to LEGAL size (8.5 x 14 inches)
    section = doc.sections[0]
    section.page_height = Inches(14)
    section.page_width = Inches(8.5)
    
    # Set 1-inch margins
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Add Title
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Content (Simple paragraph based for now, can be enhanced with html parser)
    # Removing HTML tags for basic text insertion
    clean_text = re.sub('<[^<]+?>', '', content_html)
    for line in clean_text.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.style.font.name = 'Times New Roman'
            p.style.font.size = Pt(12)
            
    doc.save(output_path)
    return output_path

def generate_tec_draft_prompt(tec_type, category, raw_input, learned_kb, indenting_member="....."):
    """
    Constructs the prompt for the LLM to draft the minutes.
    """
    patterns = learned_kb.get(tec_type, {}).get(category, learned_kb.get(tec_type, {}).get("General", {}))
    
    prompt = f"""
Draft a formal Tender Evaluation Committee (TEC) Minutes in English based on the following:

[ROLE]
You are the Member Secretary of the Tender Evaluation Committee at GSI, WR. 
Your goal is to draft a professional, legally-compliant minute for a {tec_type} evaluation.

[COMMITTEE MEMBERS]
- Chairman: {DEFAULT_COMMITTEE['chairman']}
- Member Finance: {DEFAULT_COMMITTEE['member_finance']}
- Member Secretary: {DEFAULT_COMMITTEE['member_secretary']}
- Member Indenting Division: {indenting_member}

[LEARNED PATTERNS & TONE]
- Preamble Style: {patterns.get('preamble', '')}
- Qualification Phrasing: {patterns.get('qualification_samples', '')}
- Signature Style: {patterns.get('signature_block', '')}

[VARIABLE INPUT DATA]
{raw_input}

[STANDARD DISQUALIFICATION PHRASES]
- Missing EMD: "Neither EMD submitted nor registred in relevant MSE category"
- Insufficient Experience: Use "not Compliant" / "compliant" as applicable.
- Non-submission of OEM Authorization: "No valid OEM authorisation certificate submitted."
- IP Similarity: Use this when multiple bids share the same IP address/source.
"""

    if tec_type == "Financial":
        prompt += """
[FINANCIAL TABULATION INSTRUCTIONS]
1. Extract the Item Name, Quantity, and Rates quoted by each qualified firm.
2. Present this data in a clean, professional HTML table within the "Financial Evaluation" section.
3. Rank the firms as L-1, L-2, etc., based on the total evaluated cost.
4. Ensure the table columns are clearly labeled (Firm Name, Unit Price, Tax, Total, Rank).
"""

    prompt += """
[INSTRUCTIONS]
1. Use the learned preamble style.
2. Identify the Bid ID, File Number, and Dates from the input.
3. Use the [STANDARD DISQUALIFICATION PHRASES] exactly as provided when a bidder fails on those specific grounds.
4. If this is a Financial TEC, rank the bidders as L-1, L-2, etc., as per [FINANCIAL TABULATION INSTRUCTIONS].
5. If L-1 is significantly below estimate (20%+), include a clause to seek price justification.
6. Ensure the tone is purely formal administrative English.
7. Format with sections: I. Details of Procurement, II. Salient Features, III. Technical/Financial Evaluation, IV. Recommendations.
8. End with the signature blocks for all 4 members.

Draft the complete document now.
"""
    return prompt
