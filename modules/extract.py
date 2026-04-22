"""
Noting Bot - Text Extraction Module
Handles PDF and image to text/HTML extraction using Vision LLMs.
Also handles Word (.docx) document generation from extracted content.
"""

import os
import io
import base64
import json
import logging
import hashlib
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import pytesseract
import pdfplumber
from google import genai
from google.genai import types
from modules.utils import ask_llm, CONFIG, logger, get_llm_status

# Configure Tesseract path for Windows
TESSERACT_CANDIDATES = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
]
for candidate in TESSERACT_CANDIDATES:
    if os.path.exists(candidate):
        pytesseract.pytesseract.tesseract_cmd = candidate
        logger.info(f"Using Tesseract executable at {candidate}")
        break

from modules.utils import DATA_ROOT
CACHE_DIR = DATA_ROOT / "cache"
CACHE_DIR.mkdir(parents=True, exist_ok=True)

def get_file_hash(file_path: Path) -> str:
    """Returns SHA256 hash of a file's content."""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()

def extract_text_from_file(file_path: Path = None, image_bytes: bytes = None, method: str = "standard") -> dict:
    """
    Extracts text from a PDF or Image.
    'standard' method uses local Tesseract (for images) or pdfplumber (for PDFs).
    'vision' method uses Gemini Vision.
    """
    try:
        # 1. Standard Method (Local / Hybrid)
        if method == "standard":
            if file_path and file_path.suffix.lower() == ".pdf":
                logger.info(f"Extracting text from PDF (Standard): {file_path.name}")
                text_content = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text)
                        else:
                            # Scanned page - fallback to OCR
                            logger.info(f"Page {page.page_number} appears scanned. Falling back to OCR.")
                            import fitz
                            doc = fitz.open(file_path)
                            page_fitz = doc.load_page(page.page_number - 1)
                            pix = page_fitz.get_pixmap()
                            img_data = pix.tobytes("png")
                            from PIL import Image
                            img = Image.open(io.BytesIO(img_data))
                            try:
                                text_content.append(pytesseract.image_to_string(img, lang='eng+hin'))
                            except pytesseract.TesseractNotFoundError:
                                return {
                                    "success": False,
                                    "error": "Tesseract is not installed or not found. Install Tesseract OCR and ensure it is on PATH, or set pytesseract.pytesseract.tesseract_cmd to the executable path.",
                                }
                            doc.close()
                
                return {"success": True, "text": "\n\n".join(text_content), "method": "standard_pdf"}

            # Image handling (file or bytes)
            img_data = image_bytes if image_bytes else file_path.read_bytes()
            from PIL import Image
            img = Image.open(io.BytesIO(img_data))
            logger.info("Performing Local Tesseract OCR...")
            
            # Try English + Hindi
            try:
                text = pytesseract.image_to_string(img, lang='eng+hin')
            except pytesseract.TesseractNotFoundError as ex:
                return {
                    "success": False,
                    "error": "Tesseract is not installed or not found. Install Tesseract OCR and ensure it is on PATH, or set pytesseract.pytesseract.tesseract_cmd to the correct executable path.",
                }
            except Exception:
                try:
                    text = pytesseract.image_to_string(img)
                except pytesseract.TesseractNotFoundError as ex:
                    return {
                        "success": False,
                        "error": "Tesseract is not installed or not found. Install Tesseract OCR and ensure it is on PATH, or set pytesseract.pytesseract.tesseract_cmd to the correct executable path.",
                    }

            return {"success": True, "text": text.strip(), "method": "standard_ocr"}

        # 2. Vision Method (LLM)
        else:
            api_key = CONFIG.get("gemini_api_key")
            if not api_key:
                return {"success": False, "error": "Gemini API key not configured."}

            logger.info("Vision Method: Gemini LLM configured.")
            client = genai.Client(api_key=api_key)
            
            # Use configured vision_model, fallback to gemini_model, then default
            llm_cfg = CONFIG.get("llm", {})
            model_name = llm_cfg.get("vision_model") or llm_cfg.get("gemini_model") or "gemini-2.0-flash"
            
            # Use utility normalization
            from modules.utils import _normalize_gemini_model_name
            model_name = _normalize_gemini_model_name(model_name)

            if file_path and file_path.suffix.lower() == ".pdf":
                # Handle large PDFs page-by-page with caching
                import fitz # PyMuPDF
                doc = fitz.open(file_path)
                file_hash = get_file_hash(file_path)
                total_pages = len(doc)
                logger.info(f"Vision Method: Processing {total_pages} pages for PDF: {file_path.name}")
                
                all_text_parts = []
                cache_file = CACHE_DIR / f"{file_hash}_extraction.json"
                
                # Load existing cache if available
                extracted_pages = {}
                if cache_file.exists():
                    try:
                        with open(cache_file, "r", encoding="utf-8") as f:
                            extracted_pages = json.load(f)
                        logger.info(f"Found existing cache with {len(extracted_pages)} pages.")
                    except Exception as ce:
                        logger.warning(f"Could not load cache: {ce}")

                import time
                for i in range(total_pages):
                    page_num = i + 1
                    if str(page_num) in extracted_pages:
                        logger.info(f"Page {page_num}: Using cached text.")
                        all_text_parts.append(extracted_pages[str(page_num)])
                        continue
                    
                    logger.info(f"Page {page_num}: Extracting via Vision...")
                    page = doc.load_page(i)
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # High resolution
                    img_bytes = pix.tobytes("png")
                    
                    prompt = "Transcribe all text from this document page into a structured digital format using HTML. Preserve the structural layout using <table> for tables, <b> for bold headers, and <ul> for lists. Return ONLY the HTML content without code blocks."
                    content_part = types.Part.from_bytes(data=img_bytes, mime_type="image/png")
                    
                    # Page-level retry logic
                    page_retry_count = 0
                    page_max_retries = 3
                    page_success = False
                    
                    while page_retry_count <= page_max_retries and not page_success:
                        try:
                            response = client.models.generate_content(
                                model=model_name,
                                contents=[prompt, content_part],
                                config=types.GenerateContentConfig(
                                    temperature=0.1,
                                    safety_settings=[
                                        types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                                        types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                                        types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                                        types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                                        types.SafetySetting(category="HARM_CATEGORY_CIVIC_INTEGRITY", threshold="BLOCK_NONE"),
                                    ]
                                )
                            )
                            
                            page_text = response.text.strip() if response.text else ""
                            if not page_text and response.candidates:
                                page_text = "".join([p.text for p in response.candidates[0].content.parts if hasattr(p, 'text') and p.text]).strip()
                            
                            # Clean markdown
                            if page_text.startswith("```"):
                                page_text = "\n".join(page_text.split("\n")[1:-1]).strip() if "\n" in page_text else page_text.replace("```", "")

                            all_text_parts.append(page_text)
                            extracted_pages[str(page_num)] = page_text
                            
                            # Save progress after each page
                            with open(cache_file, "w", encoding="utf-8") as f:
                                json.dump(extracted_pages, f, ensure_ascii=False, indent=2)
                            
                            page_success = True
                            # Small delay to avoid aggressive rate limiting
                            time.sleep(1)
                                
                        except Exception as pe:
                            err_msg = str(pe).lower()
                            page_retry_count += 1
                            if "429" in err_msg or "resource_exhausted" in err_msg:
                                wait_time = 2 ** page_retry_count
                                logger.warning(f"Page {page_num} Rate Limit. Retrying in {wait_time}s...")
                                time.sleep(wait_time)
                            elif page_retry_count <= page_max_retries:
                                logger.warning(f"Page {page_num} Error: {pe}. Retrying ({page_retry_count}/{page_max_retries})...")
                                time.sleep(1)
                            else:
                                logger.error(f"Page {page_num} failed after {page_max_retries} retries: {pe}")
                                all_text_parts.append(f"<p>[Error extracting page {page_num}: {str(pe)}]</p>")
                
                doc.close()
                return {
                    "success": True,
                    "text": "\n\n".join(all_text_parts),
                    "model_used": model_name,
                    "cached": True
                }

            # Single image or small file handling
            file_bytes = image_bytes if image_bytes else file_path.read_bytes()
            # Determine mime type
            mime_type = "image/png"
            if file_path:
                if file_path.suffix.lower() == ".pdf": mime_type = "application/pdf"
                elif file_path.suffix.lower() in [".jpg", ".jpeg"]: mime_type = "image/jpeg"

            prompt = "Transcribe all text from this document into a structured digital format using HTML. Preserve the structural layout using <table> for tables, <b> for bold headers, and <ul> for lists. Return ONLY the HTML content without code blocks."

            logger.info(f"Extracting single file via Vision LLM ({model_name})...")
            content_part = types.Part.from_bytes(data=file_bytes, mime_type=mime_type)
            
            # Relax safety settings to prevent false-positive blocks
            safety_settings = [
                types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
                types.SafetySetting(category="HARM_CATEGORY_CIVIC_INTEGRITY", threshold="BLOCK_NONE"),
            ]

            import time
            
            # Using a slightly more robust calling pattern with retry and fallback
            max_retries = 2
            retry_count = 0
            response = None
            
            while retry_count <= max_retries:
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=[prompt, content_part],
                        config=types.GenerateContentConfig(
                            temperature=0.1,
                            top_p=0.95,
                            safety_settings=safety_settings
                        )
                    )
                    break # Success!
                except Exception as api_err:
                    err_str = str(api_err).lower()
                    
                    # Handle 429 Resource Exhausted (Rate Limit)
                    if "429" in err_str or "resource_exhausted" in err_str or "quota" in err_str:
                        retry_count += 1
                        if retry_count <= max_retries:
                            wait_time = 2 ** retry_count
                            logger.warning(f"Gemini Rate Limit (429). Retrying in {wait_time}s... (Attempt {retry_count}/{max_retries})")
                            time.sleep(wait_time)
                            continue
                    
                    # Handle 400 (Invalid Name), 404 (Not Found), or 403 (Permission)
                    # For these (or if retries exhausted), fallback to Standard OCR immediately
                    logger.error(f"Gemini API Error: {api_err}. Falling back to Standard OCR...")
                    fallback_res = extract_text_from_file(file_path, method="standard", image_bytes=image_bytes)
                    if fallback_res.get("success"):
                        fallback_res["method"] = "standard_ocr_fallback"
                        fallback_res["fallback_reason"] = f"Gemini API Error: {str(api_err)}"
                        return fallback_res
                    else:
                        return {"success": False, "error": f"Gemini API failed ({str(api_err)}) and OCR fallback also failed: {fallback_res.get('error')}"}


            # Log safety feedback if available
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
                logger.warning(f"Gemini Prompt Feedback: {response.prompt_feedback}")

            # Check if we have candidates
            if not response.candidates or len(response.candidates) == 0:
                return {"success": False, "error": "Vision model returned no candidates. This usually happens if the content was blocked by safety filters."}

            candidate = response.candidates[0]
            if candidate.finish_reason and candidate.finish_reason != 'STOP':
                logger.warning(f"Gemini Finish Reason: {candidate.finish_reason}")
                
                # Automatic Fallback for RECITATION or SAFETY
                if candidate.finish_reason in ['SAFETY', 'RECITATION']:
                    reason_map = {
                        'SAFETY': "Safety Filters",
                        'RECITATION': "Recitation Check (Copyright)"
                    }
                    logger.info(f"Vision model blocked by {reason_map.get(candidate.finish_reason)}. Falling back to Standard OCR...")
                    fallback_res = extract_text_from_file(file_path, method="standard", image_bytes=image_bytes)
                    if fallback_res.get("success"):
                        fallback_res["method"] = "standard_ocr_fallback"
                        fallback_res["fallback_reason"] = candidate.finish_reason
                        return fallback_res
                    else:
                        return {"success": False, "error": f"Vision model blocked by {reason_map.get(candidate.finish_reason)} and OCR fallback failed: {fallback_res.get('error')}"}

            if not response.text:
                # Fallback check for parts
                parts_text = ""
                if candidate.content and candidate.content.parts:
                    parts_text = "".join([p.text for p in candidate.content.parts if hasattr(p, 'text') and p.text])
                
                if not parts_text:
                    return {"success": False, "error": "Vision model returned empty text response."}
                content = parts_text.strip()
            else:
                content = response.text.strip()

            # Clean up markdown code blocks if present
            if content.startswith("```"):
                lines = content.split("\n")
                if len(lines) > 2:
                    # Remove first line if it's ``` or ```html/markdown
                    if lines[0].startswith("```"):
                        lines = lines[1:]
                    # Remove last line if it's ```
                    if lines[-1].strip() == "```":
                        lines = lines[:-1]
                    content = "\n".join(lines).strip()
                else:
                    content = content.replace("```", "").strip()

            return {
                "success": True, 
                "text": content,
                "model_used": model_name
            }

    except Exception as e:
        logger.error(f"Extraction failed: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return {"success": False, "error": str(e)}

def generate_docx_from_html(html_content: str, output_stream: io.BytesIO):
    """
    Converts HTML/RTF content from the Quill editor into a formatted .docx file.
    """
    doc = Document()
    
    # Basic Page Setup
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Parse HTML
    soup = BeautifulSoup(html_content, "html.parser")

    def process_node(node, paragraph=None):
        if node.name is None:  # Text node
            text = node.strip()
            if text and paragraph:
                run = paragraph.add_run(text)
                return
            elif text:
                paragraph = doc.add_paragraph(text)
                return

        if node.name in ['h1', 'h2', 'h3']:
            style = 'Heading 1' if node.name == 'h1' else ('Heading 2' if node.name == 'h2' else 'Heading 3')
            doc.add_heading(node.get_text().strip(), level=int(node.name[1]))
            return

        if node.name == 'p':
            p = doc.add_paragraph()
            for child in node.children:
                if child.name == 'strong' or child.name == 'b':
                    p.add_run(child.get_text()).bold = True
                elif child.name == 'em' or child.name == 'i':
                    p.add_run(child.get_text()).italic = True
                else:
                    p.add_run(child.get_text() if child.name else str(child))
            return

        if node.name in ['ul', 'ol']:
            for li in node.find_all('li', recursive=False):
                p = doc.add_paragraph(li.get_text().strip(), style='List Bullet' if node.name == 'ul' else 'List Number')
            return
        
        # Recursive processing for other tags
        for child in node.children:
            process_node(child, paragraph)

    # Fallback to simple text if parsing is too complex or fails
    if not soup.find():
        doc.add_paragraph(html_content)
    else:
        # Better: iterate top level items
        for element in soup.contents:
            if element.name:
                if element.name == 'p':
                    p = doc.add_paragraph()
                    for r in element.children:
                        if r.name == 'strong' or r.name == 'b': p.add_run(r.get_text()).bold = True
                        elif r.name == 'em' or r.name == 'i': p.add_run(r.get_text()).italic = True
                        elif r.name == 'u': p.add_run(r.get_text()).underline = True
                        elif r.name == 'br': p.add_run("\n")
                        else: p.add_run(r.get_text() if r.name else str(r))
                elif element.name in ['h1', 'h2', 'h3', 'h4']:
                    doc.add_heading(element.get_text().strip(), level=int(element.name[1]))
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li'):
                        doc.add_paragraph(li.get_text().strip(), style='List Bullet' if element.name == 'ul' else 'List Number')
                elif element.name == 'li':
                    doc.add_paragraph(element.get_text().strip(), style='List Bullet')
                elif element.name == 'table':
                    rows = element.find_all('tr')
                    if rows:
                        col_count = max([len(r.find_all(['td', 'th'])) for r in rows], default=0)
                        if col_count > 0:
                            table = doc.add_table(rows=len(rows), cols=col_count)
                            table.style = 'Table Grid'
                            for i, row in enumerate(rows):
                                cells = row.find_all(['td', 'th'])
                                for j, cell in enumerate(cells):
                                    if j < col_count:
                                        table.cell(i, j).text = cell.get_text().strip()
                else:
                    doc.add_paragraph(element.get_text().strip())
            elif isinstance(element, str) and element.strip():
                doc.add_paragraph(element.strip())

    doc.save(output_stream)
    return output_stream


def analyze_extracted_content(text: str, context: str, file_hash: str = None) -> str:
    """Uses LLM to perform smart analysis or summarization on extracted text."""
    from modules.utils import ask_llm
    from modules.database import get_prompt_settings
    
    if not context.strip():
        context = "Summarize the document concisely while preserving all key information."

    # Check for existing summary cache if file_hash is provided
    if file_hash:
        summary_cache = CACHE_DIR / f"{file_hash}_summary.json"
        if summary_cache.exists():
            try:
                with open(summary_cache, "r", encoding="utf-8") as f:
                    cached_data = json.load(f)
                    if cached_data.get("context") == context:
                        logger.info("Using cached AI summary.")
                        return cached_data.get("processed_text")
            except Exception as e:
                logger.warning(f"Failed to read summary cache: {e}")

    prompt_settings = get_prompt_settings()
    master_prompt = prompt_settings.get("summarization_master_prompt")
    
    if master_prompt:
        try:
            full_prompt = master_prompt.format(
                user_requirement=context,
                document_text=text
            )
        except Exception as e:
            logger.warning(f"Failed to format summarization prompt: {e}")
            full_prompt = f"{master_prompt}\n\nUSER REQUIREMENT: {context}\n\nTEXT:\n{text}"
    else:
        # Fallback to hardcoded logic if DB fails
        full_prompt = f"""
        TASK: Analyze the following extracted document text.
        USER REQUIREMENT: {context}
        DOCUMENT TEXT:
        {text}
        """
    
    try:
        # If text is extremely large, Gemini 1.5/2.0 can handle it, but we might want to chunk
        # for more detailed summaries. For now, we rely on the large context window.
        result = ask_llm(full_prompt)
        
        # Save to cache if file_hash provided
        if file_hash:
            try:
                with open(summary_cache, "w", encoding="utf-8") as f:
                    json.dump({"context": context, "processed_text": result}, f, ensure_ascii=False, indent=2)
            except Exception as e:
                logger.warning(f"Failed to save summary cache: {e}")
                
        return result
    except Exception as e:
        logger.error(f"AI Analysis failed: {e}")
        raise e


def analyze_file_directly(file_path: Path = None, image_bytes: bytes = None, context: str = "", method: str = "vision") -> str:
    """
    Combines extraction and analysis in a single operation.
    Bypasses the two-step UI process for rapid summarization.
    """
    logger.info(f"Direct Analysis requested for {'file: ' + file_path.name if file_path else 'pasted image'}")
    
    ext_res = extract_text_from_file(file_path=file_path, image_bytes=image_bytes, method=method)
    if not ext_res.get("success"):
        error_msg = ext_res.get("error", "Extraction failed during direct analysis")
        logger.error(error_msg)
        raise Exception(error_msg)
    
    text = ext_res.get("text", "")
    if not text.strip():
        logger.warning("No text could be extracted for direct analysis.")
        return "No text could be extracted from the document."
        
    file_hash = get_file_hash(file_path) if file_path else None
    
    return analyze_extracted_content(text, context, file_hash=file_hash)
