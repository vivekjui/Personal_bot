"""
Noting Bot - Core Utilities
Common helper functions used across all modules.
"""

import os
import json
import logging
import re
import socket
import importlib
import time
from datetime import datetime
from pathlib import Path
from urllib import error as urllib_error
from urllib import request as urllib_request
from urllib.parse import urlsplit, unquote

import sys

# ─── Path Resolution (PyInstaller Support) ─────────────────────────────────────
if getattr(sys, 'frozen', False):
    # Running as a bundled executable
    BUNDLE_ROOT = Path(sys._MEIPASS)
else:
    # Running in a normal Python environment
    BUNDLE_ROOT = Path(__file__).parent.parent

# Persistent Data Storage (User Settings, Database, Logs)
if getattr(sys, 'frozen', False):
    # Running as a bundled executable
    if os.name == 'nt':
        DATA_ROOT = Path(os.environ.get('APPDATA', str(Path.home()))) / "APMD_Bot"
    else:
        DATA_ROOT = Path.home() / ".apmd_bot"
else:
    # Running in a normal Python environment (local dev/script mode)
    # If standard_library.json or cases.db exists in the root, use it as the data root.
    # This keeps the bot 'portable' when running from its own folder.
    project_root = Path(__file__).parent.parent
    if (project_root / "standard_library.json").exists() or (project_root / "cases.db").exists():
        DATA_ROOT = project_root
    else:
        if os.name == 'nt':
            DATA_ROOT = Path(os.environ.get('APPDATA', str(Path.home()))) / "APMD_Bot"
        else:
            DATA_ROOT = Path.home() / ".apmd_bot"

DATA_ROOT.mkdir(parents=True, exist_ok=True)
BOT_ROOT = DATA_ROOT # For backward compatibility in other modules

# All databases moved to a dedicated subfolder if using AppData, or root if portable
if DATA_ROOT != BUNDLE_ROOT and not getattr(sys, 'frozen', False):
    DB_ROOT = DATA_ROOT / "db"
else:
    DB_ROOT = DATA_ROOT

DB_ROOT.mkdir(parents=True, exist_ok=True)

CONFIG_PATH = DATA_ROOT / "config.json"
LEGACY_LLM_PROMPT_VALUES = {}
DEFAULT_GEMINI_MODEL = "gemini-2.0-flash"


def _normalize_gemini_model_name(model_name: str) -> str:
    """Normalize Gemini model identifiers to the new Gemini SDK naming conventions."""
    if not model_name:
        model_name = DEFAULT_GEMINI_MODEL
    
    # Normalize input: lowercase and replace spaces/underscores with hyphens
    model_id = model_name.lower().strip().replace(" ", "-").replace("_", "-")
    
    # Remove models/ prefix if present for logic comparison
    if model_id.startswith("models/"):
        model_id = model_id.split("models/", 1)[1]
    
    # Mapping for specific model identifiers
    mappings = {
        "gemma3-27b": "gemma-3-27b-it",
        "gemini-3.1-flash-lite": "gemini-3.1-flash-lite-preview",
        "gemini-1.5-flash": "gemini-2.0-flash",
        "gemini-1.5-flash-latest": "gemini-2.0-flash",
        "gemini-1.5-pro": "gemini-2.5-pro",
        "gemini-flash-latest": "gemini-2.0-flash",
        "gemini-3-flash": "gemini-2.0-flash", # Fallback for unsupported Gemini 3 Flash
    }
    
    if model_id in mappings:
        return mappings[model_id]
        
    return model_id


def _default_config() -> dict:
    """Return a safe default configuration for first-run bootstrap."""
    return {
        "gemini_api_key": "",
        "paths": {
            "cases_dir": "Cases",
            "templates_dir": "templates",
            "criteria_file": "criteria/bid_evaluation_criteria.json",
            "database": str(DB_ROOT / "cases.db"),
            "logs": str(DATA_ROOT / "logs"),
        },
        "reminders": {
            "emd_alert_days_before": 30,
            "ps_alert_days_before": 30,
            "completion_alert_days_before": 15,
            "dlp_alert_days_before": 15,
        },
        "portals": {
            "gem": "https://gem.gov.in",
            "cppp": "https://eprocure.gov.in",
        },
        "dashboard": {
            "host": "127.0.0.1",
            "port": 5006,
            "enable_widget": False,
            "debug": False,
        },
        "llm": {
            "provider": "gemma3_27b",
            "gemini_model": DEFAULT_GEMINI_MODEL,
            "temperature": 0.3,
            "context_length": 8192,
            "timeout_seconds": 120,
        },
        "rag": {
            "enabled": True,
            "chunk_size": 800,
            "chunk_overlap": 150,
            "top_k_results": 5,
            "min_relevance_pct": 40,
            "kb_dir": str(DATA_ROOT / "knowledge_base"),
            "embedding_model": "all-MiniLM-L6-v2",
        },
        "network": {
            "proxy_mode": "manual",
            "proxy_server": "http://10.6.0.9",
            "proxy_port": "3128",
            "proxy_username": "",
            "proxy_password": "",
        },
    }


def ensure_bundle_resource(rel_path: str) -> Path:
    """
    Copy a bundled file or directory into user-writable app storage on first run.
    Returns the writable target path.
    """
    import shutil

    source = BUNDLE_ROOT / rel_path
    target = DATA_ROOT / rel_path

    if target.exists():
        return target

    if source.exists():
        target.parent.mkdir(parents=True, exist_ok=True)
        if source.is_dir():
            shutil.copytree(source, target, dirs_exist_ok=True)
        else:
            shutil.copy2(source, target)

    return target


STANDARD_LIBRARY_PATH = ensure_bundle_resource("standard_library.json")
PROCUREMENT_STAGES_PATH = ensure_bundle_resource("procurement_stages.json")
DEFAULT_KB_DIR = ensure_bundle_resource("knowledge_base")

# ─── Logging Setup ─────────────────────────────────────────────────────────────
# Pre-define a basic logger for use during config loading
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s"
)
logger = logging.getLogger("Noting_Bot")

DEFAULT_NOTING_MASTER_PROMPT = """You are an expert procurement professional.

Draft an official noting in Hindi by default. Convert Hinglish into proper official Hindi.
If any sentence is in English, convert it to Hindi unless the source content must stay as-is. Use the available reference context and writing style examples when helpful.
-Table data also to be converted in Hindi.
-बोली to be replaced with निविदा
- use smart intelligence to Ensure the firm name / contract name etc remain same throughout if the user forget to update in later paragraph / content the name of firm / contract number etc.
- बोलीदाता to be replaced with निविदाकर्ता
- Use English alternative (in bracket) of complex hindi word / terminology
- If a highly relevant draft or template is found in the Preferred Style Examples, you MUST follow its exact structure, tone, and phrasing, only substituting the specific details from the additional context provided.
- Always use Markdown tables for any data comparisons, price lists, or tabular reports.

Additional Context:
{additional_context}

Reference Context:
{rag_context}

Preferred Style Examples:
{user_style_examples}
Check if the first paragraph modified by the user contains firm name as "x" and forget to replace in subsequent paragraph, then correct this. Check for calculations made (correct if wrong calculated). If there is any Figure in Rupees, then same may be written in word in bracket also.
Check for instruction in additional context. rearrange the noting text as per context. Add contextual topic in appropriate place. Return only the final noting text without subject or sub-heading.
"""

DEFAULT_EMAIL_MASTER_PROMPT = """You are an expert Indian Government official drafting a formal email.

Refine the provided draft into a polished official email body in {target_language}.
- Keep the output as an email, not a file noting.
- Never add the closing line "\u092b\u093e\u0907\u0932 \u0906\u092a\u0915\u0947 \u0905\u0935\u0932\u094b\u0915\u0928\u093e\u0930\u094d\u0925 \u092a\u094d\u0930\u0938\u094d\u0924\u0941\u0924 \u0939\u0948 \u0964" or any similar file-submission line unless the user explicitly asks for it.
- If the draft already contains a closing/sign-off, keep only one appropriate closing and do not repeat it.
- Preserve names, references, numbers, contract details, and email-specific structure unless the user asks to change them.
- Follow the user's stored style and learned wording preferences whenever relevant.

Draft Content:
{draft_content}

Additional Instructions:
{additional_instructions}

Preferred Style Examples:
{user_style_examples}

Style Summary:
{style_summary}

Learning Instructions:
{learning_instructions}

Return only the final email content without explanation.
"""

DEFAULT_SUMMARIZATION_MASTER_PROMPT = """Analyze the following extracted document text based on the USER REQUIREMENT.

USER REQUIREMENT: {user_requirement}

GUIDELINES:
1. Provide a structured, professional summary or analysis as per the user requirement.
2. Maintain an official, government-standard tone.
3. Highlight key dates, entities (firms, individuals), monetary amounts, and action items.
4. If technical evaluation is involved, clearly list qualification status for each vendor.
5. Use Markdown tables or bullet points for clarity.
6. Provide the result in clean, well-formatted Rich Text (HTML).

EXTRACTED TEXT:
---
{document_text}
---
"""

DEFAULT_KNOWHOW_MASTER_PROMPT = """You are an expert Government Official and Procurement Specialist.
Your task is to answer user questions based STRICTLY on the provided Knowledge Base context.

If the information is not in the context, say you don't know rather than hallucinating.
Always provide rule numbers or circular references if mentioned in the context.

ANSWER PATTERN (strictly follow this order):
1. GFR 2017: Relevant clause and description (if found in context).
2. Manual for Procurement of Goods: Relevant clause and description (if found in context).
3. GeM ATC (Additional Terms & Conditions): Relevant clause and description (if found in context).
4. GSI Manual: Relevant clause and description (if found in context).
5. Web Search Result / Supplemental Info: Provide relevant external or supplemental info.
6. Advisory: Provide a practical advisory or recommendation for the user.

=== LEARNING CONTEXT ===
{learning_context}

=== KNOWLEDGE BASE CONTEXT ===
{context}
==============================

User Question: {prompt}

Provide a helpful, precise answer in {target_language}.
"""
DEFAULT_QA_SYSTEM_PROMPT = DEFAULT_KNOWHOW_MASTER_PROMPT

def load_config() -> dict:
    """Load the main config.json file. If missing in DATA_ROOT, copy from BUNDLE_ROOT.

    When a configuration is first created or read, ensure the network/proxy
    section contains sane defaults.  The application should always attempt to
    route traffic through the corporate proxy at 10.6.0.9:3128 unless the user
    deliberately turns the proxy off.  This function will insert those defaults
    and persist them back to disk if necessary.
    """
    global LEGACY_LLM_PROMPT_VALUES
    import shutil
    wrote_back = False

    if not CONFIG_PATH.exists():
        bootstrap_candidates = [
            BUNDLE_ROOT / "config.json",
            BUNDLE_ROOT / "config.example.json",
        ]
        copied = False
        for candidate in bootstrap_candidates:
            if candidate.exists():
                shutil.copy2(candidate, CONFIG_PATH)
                copied = True
                break
        if not copied:
            with open(CONFIG_PATH, "w", encoding="utf-8") as f:
                json.dump(_default_config(), f, indent=2)

    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        cfg = json.load(f)

    if not isinstance(cfg, dict):
        cfg = _default_config()
        wrote_back = True

    # Ensure all master prompts are in config
    prompts = cfg.setdefault("prompts", {})
    if "noting_master" not in prompts:
        prompts["noting_master"] = DEFAULT_NOTING_MASTER_PROMPT
        wrote_back = True
    if "email_master" not in prompts:
        prompts["email_master"] = DEFAULT_EMAIL_MASTER_PROMPT
        wrote_back = True
    if "knowhow_master" not in prompts:
        prompts["knowhow_master"] = DEFAULT_KNOWHOW_MASTER_PROMPT
        wrote_back = True

    cfg.setdefault("gemini_api_key", "")

    paths = cfg.setdefault("paths", {})
    path_defaults = {
        "cases_dir": "Cases",
        "templates_dir": "templates",
        "criteria_file": "criteria/bid_evaluation_criteria.json",
        "database": str(DATA_ROOT / "cases.db"),
        "logs": str(DATA_ROOT / "logs"),
    }
    for key, value in path_defaults.items():
        if not paths.get(key):
            paths[key] = value
            wrote_back = True

    dashboard = cfg.setdefault("dashboard", {})
    dashboard_defaults = {
        "host": "127.0.0.1",
        "port": 5006,
        "enable_widget": False,
        "debug": False,
    }
    for key, value in dashboard_defaults.items():
        if key not in dashboard:
            dashboard[key] = value
            wrote_back = True

    reminders = cfg.setdefault("reminders", {})
    reminder_defaults = {
        "emd_alert_days_before": 30,
        "ps_alert_days_before": 30,
        "completion_alert_days_before": 15,
        "dlp_alert_days_before": 15,
    }
    for key, value in reminder_defaults.items():
        if key not in reminders:
            reminders[key] = value
            wrote_back = True

    portals = cfg.setdefault("portals", {})
    portal_defaults = {
        "gem": "https://gem.gov.in",
        "cppp": "https://eprocure.gov.in",
    }
    for key, value in portal_defaults.items():
        if key not in portals:
            portals[key] = value
            wrote_back = True

    # ensure network defaults for proxy
    net = cfg.setdefault("network", {})
    # corporate proxy defaults
    DEFAULT_PROXY_SERVER = "http://10.6.0.9"
    DEFAULT_PROXY_PORT = "3128"

    # if the existing configuration explicitly disabled proxy but there is
    # no server/port defined, treat this as an uninitialised state and fall
    # back to our default manual proxy.  This prevents the situation where
    # the file imported from AppData had "off" carried over from an old
    # install and the bot therefore never used the proxy.
    if (net.get("proxy_mode") == "off" and
        not net.get("proxy_server") and
        not net.get("proxy_port")):
        net["proxy_mode"] = "manual"
        net["proxy_server"] = DEFAULT_PROXY_SERVER
        net["proxy_port"] = DEFAULT_PROXY_PORT
        wrote_back = True

    # if no mode provided at all assume manual
    if not net.get("proxy_mode"):
        net["proxy_mode"] = "manual"
        wrote_back = True
    # if mode is manual ensure server/port exist
    if net.get("proxy_mode") == "manual":
        if not net.get("proxy_server"):
            net["proxy_server"] = DEFAULT_PROXY_SERVER
            wrote_back = True
        if not net.get("proxy_port"):
            net["proxy_port"] = DEFAULT_PROXY_PORT
            wrote_back = True
    # fill out username/password keys if missing
    for key in ("proxy_username", "proxy_password"):
        if key not in net:
            net[key] = ""
            wrote_back = True

    # ensure default LLM settings (Gemma 3 27B provider)
    llm = cfg.setdefault("llm", {})
    if not llm.get("provider"):
        llm["provider"] = "gemma3_27b"
        wrote_back = True
    # model fallback and other parameters
    if "gemini_model" not in llm:
        llm["gemini_model"] = DEFAULT_GEMINI_MODEL
        wrote_back = True
    else:
        normalized_model = _normalize_gemini_model_name(llm.get("gemini_model"))
        if normalized_model != llm.get("gemini_model"):
            llm["gemini_model"] = normalized_model
            wrote_back = True
    if "temperature" not in llm:
        llm["temperature"] = 0.3
        wrote_back = True
    if "context_length" not in llm:
        llm["context_length"] = 8192
        wrote_back = True
    if "timeout_seconds" not in llm:
        llm["timeout_seconds"] = 120
        wrote_back = True
    LEGACY_LLM_PROMPT_VALUES = {}
    for prompt_key in ("noting_master_prompt", "email_master_prompt", "qa_system_prompt"):
        legacy_value = llm.pop(prompt_key, None)
        if legacy_value:
            LEGACY_LLM_PROMPT_VALUES[prompt_key] = legacy_value
            wrote_back = True

    if wrote_back:
        try:
            with open(CONFIG_PATH, "w", encoding="utf-8") as outf:
                json.dump(cfg, outf, indent=2)
        except Exception:
            # if writing fails just ignore; we still use values in memory
            pass
    
    # ── Path Sanitization (Portability) ────────────────────────────────────────
    # If paths are hardcoded to a different machine/folder, re-base them to DATA_ROOT
    if "paths" in cfg:
        for key, val in cfg["paths"].items():
            if isinstance(val, str):
                p = Path(val)
                # Keep persistent runtime files in user data even if the template
                # config uses relative placeholders from the repository.
                if key == "database" and not p.is_absolute():
                    cfg["paths"][key] = str(DATA_ROOT / p.name)
                    wrote_back = True
                    continue
                if key == "logs" and not p.is_absolute():
                    cfg["paths"][key] = str(DATA_ROOT / p.name)
                    wrote_back = True
                    continue
                # If path is absolute and doesn't exist, OR it contains the old dev folder name
                if (p.is_absolute() and not p.exists()) or "APMD_eOffice_Bot" in val:
                    basename = p.name
                    if key == "database":
                        cfg["paths"][key] = str(DATA_ROOT / "cases.db")
                    elif key == "logs":
                        cfg["paths"][key] = str(DATA_ROOT / "logs")
                    else:
                        cfg["paths"][key] = str(DATA_ROOT / basename)
    
    if "rag" in cfg and "kb_dir" in cfg["rag"]:
        val = cfg["rag"]["kb_dir"]
        p = Path(val)
        if not p.is_absolute():
            cfg["rag"]["kb_dir"] = str(DATA_ROOT / p.name)
            wrote_back = True
        elif (p.is_absolute() and not p.exists()) or "APMD_eOffice_Bot" in val:
            cfg["rag"]["kb_dir"] = str(DATA_ROOT / "knowledge_base")
    else:
        rag = cfg.setdefault("rag", {})
        rag_defaults = {
            "enabled": True,
            "chunk_size": 800,
            "chunk_overlap": 150,
            "top_k_results": 5,
            "min_relevance_pct": 40,
            "kb_dir": str(DATA_ROOT / "knowledge_base"),
            "embedding_model": "all-MiniLM-L6-v2",
        }
        for key, value in rag_defaults.items():
            if key not in rag:
                rag[key] = value
                wrote_back = True

    # ── Database Migration ─────────────────────────────────────────────────────
    if "paths" in cfg and "database" in cfg["paths"]:
        import shutil
        db_path = Path(cfg["paths"]["database"])
        if not db_path.exists():
            # Try to find it in BUNDLE_ROOT (legacy location)
            legacy_db = BUNDLE_ROOT / "cases.db"
            if legacy_db.exists():
                db_path.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(legacy_db, db_path)
                logger.info(f"Migrated database from legacy location to {db_path}")

    return cfg

CONFIG = load_config()

def _ensure_localhost_no_proxy() -> None:
    """Always bypass proxies for local Selenium and DevTools traffic."""
    required_hosts = ["127.0.0.1", "localhost"]

    existing = []
    for key in ("NO_PROXY", "no_proxy"):
        value = os.environ.get(key, "")
        if value:
            existing.extend([part.strip() for part in value.split(",") if part.strip()])

    merged = []
    seen = set()
    for host in existing + required_hosts:
        host_l = host.lower()
        if host_l not in seen:
            seen.add(host_l)
            merged.append(host)

    merged_value = ",".join(merged)
    os.environ["NO_PROXY"] = merged_value
    os.environ["no_proxy"] = merged_value


def apply_proxy_settings():
    """Apply proxy settings from config to the process environment."""
    net = CONFIG.get("network", {})
    mode = net.get("proxy_mode", "off")
    
    if mode == "off":
        # Force clear proxy env vars
        keys = ["HTTP_PROXY", "HTTPS_PROXY", "http_proxy", "https_proxy"]
        for k in keys:
            if k in os.environ:
                del os.environ[k]
        _ensure_localhost_no_proxy()
        logger.debug("Proxy: Off")
    elif mode == "manual":
        server = net.get("proxy_server", "").strip()
        port = net.get("proxy_port", "").strip()
        user = net.get("proxy_username", "").strip()
        pw = net.get("proxy_password", "").strip()
        
        if server and port:
            # strip any scheme from server, we'll add http:// ourselves
            if server.startswith("http://"):
                server_clean = server[len("http://"):]
            elif server.startswith("https://"):
                server_clean = server[len("https://"):]
            else:
                server_clean = server

            # Construct proxy URL
            if user and pw:
                proxy_url = f"http://{user}:{pw}@{server_clean}:{port}"
            else:
                proxy_url = f"http://{server_clean}:{port}"
            
            os.environ["HTTP_PROXY"] = proxy_url
            os.environ["HTTPS_PROXY"] = proxy_url
            os.environ["http_proxy"] = proxy_url
            os.environ["https_proxy"] = proxy_url
            _ensure_localhost_no_proxy()
            logger.info(f"Proxy: Manual ({server_clean}:{port})")
    elif mode == "system":
        # Python's requests/httpx/urllib typically auto-detect system proxy
        # if environment variables are NOT set.
        keys = ["HTTP_PROXY", "HTTPS_PROXY", "http_proxy", "https_proxy"]
        for k in keys:
            if k in os.environ:
                del os.environ[k]
        _ensure_localhost_no_proxy()
        logger.info("Proxy: System (Auto-detecting)")

# Apply settings immediately on import
apply_proxy_settings()


def get_requests_proxies():
    """Return a proxy dictionary suitable for ``requests`` or similar libraries.

    The values are derived from the current configuration; if manual mode is
    active the returned dict will include both http and https entries pointing
    at the configured server/port.  An empty dict is returned if no proxy is
    required.
    """
    net = CONFIG.get("network", {})
    if net.get("proxy_mode") == "manual":
        server = net.get("proxy_server", "").strip()
        port = net.get("proxy_port", "").strip()
        if server and port:
            # strip any leading scheme
            if server.startswith("http://"):
                server_clean = server[len("http://"):]
            elif server.startswith("https://"):
                server_clean = server[len("https://"):]
            else:
                server_clean = server
            proxy_url = f"http://{server_clean}:{port}"
            return {"http": proxy_url, "https": proxy_url}
    return {}


def test_proxy_connection(url: str = "http://httpbin.org/ip", timeout: int = 5) -> dict:
    """Quickly verify that the proxy is being used by doing a simple HTTP call.

    Returns a dictionary containing the response JSON (if any) along with the
    proxy settings that were applied.  Errors are raised as exceptions so the
    caller can handle connectivity problems.
    """
    import requests
    proxies = get_requests_proxies()
    resp = requests.get(url, timeout=timeout, proxies=proxies)
    return {"status_code": resp.status_code, "json": resp.json(), "proxies": proxies}

# ─── Final Logging Configuration (with File Handler) ───────────────────────────
log_dir = Path(CONFIG["paths"]["logs"])
log_dir.mkdir(parents=True, exist_ok=True)

for handler in logger.handlers[:]:
    logger.removeHandler(handler)

file_handler = logging.FileHandler(log_dir / "bot_activity.log", encoding="utf-8")
file_handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(name)s: %(message)s"))
logger.addHandler(file_handler)
logger.addHandler(logging.StreamHandler())


# ─── Folder Utilities ──────────────────────────────────────────────────────────
def get_case_folder(case_id: str, subfolder: str = "") -> Path:
    """Return the path for a specific case folder, creating it if needed."""
    cases_dir = Path(CONFIG["paths"]["cases_dir"])
    case_path = cases_dir / case_id
    if subfolder:
        case_path = case_path / subfolder
    case_path.mkdir(parents=True, exist_ok=True)
    return case_path


def ensure_folder(path: str) -> Path:
    """Ensure a folder exists, creating it if necessary."""
    p = Path(path)
    p.mkdir(parents=True, exist_ok=True)
    return p


# ─── Text & Date Utilities ─────────────────────────────────────────────────────
def sanitize_filename(name: str) -> str:
    """Remove characters unsafe for filenames."""
    return re.sub(r'[<>:"/\\|?*]', '_', name).strip()


def today_str() -> str:
    """Return today's date in DD-MM-YYYY format."""
    return datetime.now().strftime("%d-%m-%Y")


def parse_date_flexible(date_str: str):
    """
    Try to parse a date string in multiple common Indian formats.
    Returns a datetime object or None if parsing fails.
    """
    from dateutil import parser as date_parser
    try:
        return date_parser.parse(date_str, dayfirst=True)
    except Exception:
        return None


def days_until(target_date) -> int:
    """Return number of days from today until target_date."""
    if isinstance(target_date, str):
        target_date = parse_date_flexible(target_date)
    if target_date is None:
        return None
    delta = target_date.date() - datetime.now().date()
    return delta.days


# ─── LLM Router (Ollama primary → Gemini fallback) ────────────────────────────
# Ollama runs locally via `ollama run phi4-mini` (no API key needed).
# Gemini is used as fallback if Ollama is not running or not configured.

# --- Local Ollama Integration Removed ---





def _ask_gemini_direct(prompt: str, override_model: str = None) -> str:
    """Send prompt to Google Gemini API (or Gemma via API) and return response text."""
    from google import genai
    api_key = CONFIG.get("gemini_api_key", "")
    if not api_key or api_key == "YOUR_GEMINI_API_KEY_HERE":
        raise ValueError("Gemini API key not configured in config.json.")
    
    model_name = override_model if override_model else CONFIG.get("llm", {}).get("gemini_model", DEFAULT_GEMINI_MODEL)
    model_name = _normalize_gemini_model_name(model_name)
    
    client = genai.Client(api_key=api_key)
    response = client.models.generate_content(
        model=f"models/{model_name}", # Add 'models/' prefix for API
        contents=prompt
    )
    try:
        if response.text:
            return response.text.strip()
        return "[No text returned from Gemini]"
    except Exception:
        # Check if it was blocked
        if hasattr(response, 'candidates') and response.candidates:
            cand = response.candidates[0]
            if cand.finish_reason:
                return f"[Gemini response blocked. Reason: {cand.finish_reason}]"
        return f"[Error accessing Gemini response text]"


def _ask_groq_direct(prompt: str, override_model: str = None) -> str:
    """Send prompt to Groq API using OpenAI-compatible client."""
    import requests
    api_key = CONFIG.get("llm", {}).get("groq_api_key", "")
    if not api_key:
        raise ValueError("Groq API key not configured.")
    
    model_name = override_model if override_model else CONFIG.get("llm", {}).get("groq_model", "llama-3.3-70b-versatile")
    
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model_name,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": CONFIG.get("llm", {}).get("temperature", 0.3)
    }
    
    from modules.utils import get_requests_proxies
    response = requests.post(url, headers=headers, json=payload, timeout=60, proxies=get_requests_proxies())
    response.raise_for_status()
    data = response.json()
    return data["choices"][0]["message"]["content"].strip()


def ask_llm(prompt: str, context: str = "") -> str:
    """
    Universal LLM call with customizable priority chains based on config.
    """
    full_prompt = f"{context}\n\n{prompt}" if context else prompt
    llm_cfg     = CONFIG.get("llm", {})
    provider    = llm_cfg.get("provider", "gemini")

    def try_gemini():
        logger.debug("LLM: trying Gemini")
        return _ask_gemini_direct(full_prompt, override_model=llm_cfg.get("gemini_model"))

    def try_groq():
        logger.debug("LLM: trying Groq")
        return _ask_groq_direct(full_prompt, override_model=llm_cfg.get("groq_model"))

    def try_gemma():
        logger.debug("LLM: trying Gemma 3")
        return _ask_gemini_direct(full_prompt, override_model="gemma3_27b")
        
    if provider == "groq":
        try: return try_groq()
        except Exception as e:
            logger.warning(f"Groq failed ({e}). Falling back to Gemini.")
            return try_gemini()
    elif provider == "gemma3_27b":
        try: return try_gemma()
        except Exception as e1:
            logger.warning(f"Gemma 3 failed ({e1}). Falling back to Gemini.")
            try: return try_gemini()
            except Exception as e2:
                return f"[AI Error: Gemini backends failed]"
    else: # "gemini" or "auto"
        try: return try_gemini()
        except Exception as e1:
            try: return try_gemma()
            except Exception as e2:
                return f"[AI Error: Gemini backends failed]"


# Backward-compatible alias — all existing modules call ask_gemini(); no changes needed.
def ask_gemini(prompt: str, context: str = "") -> str:
    return ask_llm(prompt, context)


def get_llm_status() -> dict:
    """Return status information about the LLM configuration."""
    llm_cfg = CONFIG.get("llm", {})
    provider = llm_cfg.get("provider", "gemini")
    return {
        "provider": provider,
        "active_backend": provider.capitalize() if provider else "None",
        "gemini_model": llm_cfg.get("gemini_model", DEFAULT_GEMINI_MODEL),
        "groq_model": llm_cfg.get("groq_model", "llama-3.3-70b-versatile"),
        "temperature": llm_cfg.get("temperature", 0.3),
        "context_length": llm_cfg.get("context_length", 8192),
        "gemini_key_set": bool(CONFIG.get("gemini_api_key")),
        "has_groq_key": bool(llm_cfg.get("groq_api_key"))
    }


def list_available_models() -> dict:
    """Fetch available models from configured providers and filter for suitability."""
    results = {"gemini": [], "groq": []}
    
    # Define keywords/prefixes for models suitable for general LLM work and Vision tasks
    SUITABLE_PREFIXES = [
        "gemini-1.5", "gemini-2.0", "gemini-1.0-pro", "gemini-pro-vision",
        "llama-3.1", "llama-3.2", "llama-3.3", "llama-3-", "llama3-",
        "gemma-2", "gemma2", "gemma3", "mixtral-8x7b", "llama-3.2-11b-vision",
        "llama-3.2-90b-vision"
    ]

    def is_suitable(model_id):
        mid = model_id.lower()
        if any(p in mid for p in SUITABLE_PREFIXES):
            # Exclude guard and whisper models explicitly
            if any(evil in mid for evil in ["guard", "whisper", "distil"]):
                return False
            return True
        return False

    # Google Gemini Models
    try:
        from google import genai
        api_key = CONFIG.get("gemini_api_key", "")
        if api_key and api_key != "YOUR_GEMINI_API_KEY_HERE":
            client = genai.Client(api_key=api_key)
            for m in client.models.list():
                model_id = getattr(m, 'name', None)
                if not model_id:
                    continue
                model_id = model_id.replace("models/", "")
                if is_suitable(model_id):
                    results["gemini"].append({
                        "id": model_id,
                        "name": getattr(m, 'display_name', model_id),
                        "description": getattr(m, 'description', "")
                    })
    except Exception as e:
        logger.warning(f"Failed to list Gemini models: {e}")

    # Groq Models
    try:
        api_key = CONFIG.get("llm", {}).get("groq_api_key", "")
        if api_key:
            import requests
            from modules.utils import get_requests_proxies
            url = "https://api.groq.com/openai/v1/models"
            headers = {"Authorization": f"Bearer {api_key}"}
            r = requests.get(url, headers=headers, timeout=5, proxies=get_requests_proxies())
            if r.status_code == 200:
                data = r.json()
                for m in data.get("data", []):
                    model_id = m["id"]
                    if is_suitable(model_id):
                        results["groq"].append({
                            "id": model_id,
                            "name": model_id,
                            "description": f"Provider: {m.get('owned_by', 'Groq')}"
                        })
    except Exception as e:
        logger.warning(f"Failed to list Groq models: {e}")

    return results


def extract_text_custom(file_path: str = None, image_bytes: bytes = None, method: str = "standard") -> str:
    """
    Integrated extraction for PDFs or Images.
    If standard method fails on image, it returns empty.
    If vision method is used, it uses Gemini Flash.
    """
    if file_path:
        ext = os.path.splitext(file_path.lower())[1]
        if ext == ".pdf":
            if method == "vision":
                return extract_text_from_pdf_vision(file_path)
            else:
                return extract_text_from_pdf(file_path)
        elif ext in [".png", ".jpg", ".jpeg"]:
            with open(file_path, "rb") as f:
                image_bytes = f.read()
    
    if image_bytes:
        return extract_text_from_image(image_bytes, method=method)
    
    return ""
def extract_text_from_pdf(pdf_path: str) -> str:
    """Compatibility wrapper for rag_engine and other modules."""
    from modules.extract import extract_text_from_file
    res = extract_text_from_file(file_path=Path(pdf_path), method="standard")
    return res.get("text", "")



def extract_text_from_pdf_vision(pdf_path: str) -> str:
    """Force Vision extraction for all pages of a PDF."""
    vision_text = []
    gemini_key = CONFIG.get("gemini_api_key", "")
    if not gemini_key or gemini_key == "YOUR_GEMINI_API_KEY_HERE":
        return "[Error: No Gemini API Key for Vision]"

    try:
        import fitz  # PyMuPDF
        from google import genai
        from google.genai import types
        import tempfile
        from PIL import Image

        client = genai.Client(api_key=gemini_key)
        model_name = _normalize_gemini_model_name(CONFIG.get("llm", {}).get("gemini_model", DEFAULT_GEMINI_MODEL))

        doc = fitz.open(pdf_path)
        for i in range(len(doc)):
            page = doc.load_page(i)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_bytes = pix.tobytes("png")
            content_part = types.Part.from_bytes(data=img_bytes, mime_type="image/png")
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=[
                        "Extract all text from this page exactly as it appears. Preserve layout using Markdown tables if needed. Return ONLY the text.",
                        content_part
                    ],
                    config=types.GenerateContentConfig(temperature=0.1)
                )
                if response and response.text:
                    vision_text.append(response.text)
                elif response and hasattr(response, 'candidates') and response.candidates:
                    # Fallback for empty text property but has content
                    parts_text = "".join([p.text for p in response.candidates[0].content.parts if hasattr(p, 'text') and p.text])
                    if parts_text:
                        vision_text.append(parts_text)
            except Exception as page_err:
                logger.warning(f"Failed to extract text from page {i}: {page_err}")
                continue
        doc.close()
    except Exception as e:
        logger.error(f"PDF Vision extraction failed: {e}")
        return f"[Error: {e}]"

    return "\n\n".join(vision_text)


def extract_text_from_image(image_bytes: bytes, method: str = "standard") -> str:
    """
    Extract text from an image. 
    'standard' mode uses local Tesseract OCR.
    'vision' mode uses Gemini Vision via the configured model.
    """
    import io
    from PIL import Image
    
    try:
        img = Image.open(io.BytesIO(image_bytes))
        
        if method == "standard":
            import pytesseract
            # Optional: Specify tesseract path if needed, though usually on path in this environment
            # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            logger.info("Performing Standard OCR via Pytesseract...")
            
            # Try English + Hindi if available, else default
            try:
                text = pytesseract.image_to_string(img, lang='eng+hin')
            except:
                text = pytesseract.image_to_string(img)
            return text.strip()
            
        else: # method == "vision"
            gemini_key = CONFIG.get("gemini_api_key", "")
            if not gemini_key or gemini_key == "YOUR_GEMINI_API_KEY_HERE":
                return "[Error: No Gemini API Key for Vision extraction]"
                
            from google import genai
            from google.genai import types
            client = genai.Client(api_key=gemini_key)
            model_name = _normalize_gemini_model_name(CONFIG.get("llm", {}).get("gemini_model", DEFAULT_GEMINI_MODEL))
            
            logger.info("Performing Vision LLM extraction via Gemini...")
            prompt = "Extract all text from this image. Format it cleanly, preserving tables or complex layouts if they exist. Return ONLY the extracted text."
            image_part = types.Part.from_bytes(data=image_bytes, mime_type='image/png')
            response = client.models.generate_content(
                model=model_name,
                contents=[prompt, image_part]
            )
            return response.text.strip()
            
    except Exception as e:
        logger.error(f"Image extraction failed ({method}): {e}")
        return f"[Extraction Error: {e}]"


# ─── DOCX Utilities ────────────────────────────────────────────────────────────
def fill_docx_template(template_path: str, replacements: dict, output_path: str) -> str:
    """
    Fill a .docx template by replacing placeholder text with actual values.
    Placeholders in template should be in format: {{PLACEHOLDER_NAME}}
    Returns the output path on success.
    """
    from docx import Document
    doc = Document(template_path)

    def replace_in_paragraph(para):
        for key, value in replacements.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in para.text:
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))

    for para in doc.paragraphs:
        replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    doc.save(output_path)
    logger.info(f"Document saved: {output_path}")
    return output_path


def create_docx_from_html(html_content: str, output_path: str, title: str = "") -> str:
    """
    Create a .docx file from HTML content (from Quill/RTF).
    Handles basic tags: b, i, u, p, h1, h2, h3, ul, li, table.
    """
    from docx import Document
    from docx.shared import Pt
    from bs4 import BeautifulSoup

    doc = Document()
    soup = BeautifulSoup(html_content, 'html.parser')

    if title:
        doc.add_heading(title, level=0)

    for element in soup.contents:
        if element.name in ['h1', 'h2', 'h3']:
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level=level)
        elif element.name == 'p':
            p = doc.add_paragraph()
            _process_html_children(element, p)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                p = doc.add_paragraph(style='List Bullet')
                _process_html_children(li, p)
        elif element.name == 'ol':
            for li in element.find_all('li'):
                p = doc.add_paragraph(style='List Number')
                _process_html_children(li, p)
        elif element.name == 'table':
            rows = element.find_all('tr')
            if not rows: continue
            first_row = rows[0].find_all(['td', 'th'])
            table = doc.add_table(rows=len(rows), cols=len(first_row))
            table.style = 'Table Grid'
            for r_idx, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for c_idx, cell in enumerate(cells):
                    if c_idx < len(first_row):
                        _process_html_children(cell, table.cell(r_idx, c_idx).paragraphs[0])
        elif element.name is None: # text node
            if element.strip():
                doc.add_paragraph(element)

    doc.save(output_path)
    return output_path

def _process_html_children(element, paragraph):
    """Recursively process HTML children into a docx paragraph's runs."""
    from bs4 import NavigableString
    for child in element.children:
        if isinstance(child, NavigableString):
            paragraph.add_run(str(child))
        else:
            run = paragraph.add_run(child.get_text())
            if child.name == 'b' or child.name == 'strong':
                run.bold = True
            if child.name == 'i' or child.name == 'em':
                run.italic = True
            if child.name == 'u':
                run.underline = True
            if child.name == 'br':
                paragraph.add_run('\n')

def create_docx_from_text(content: str, output_path: str, title: str = "") -> str:
    """Fallback if only text is available."""
    from docx import Document
    doc = Document()
    if title: doc.add_heading(title, 0)
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)
    return output_path


def find_free_port(start_port: int, max_tries: int = 10) -> int:
    """Find an available port starting from start_port."""
    port = start_port
    for _ in range(max_tries):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            if s.connect_ex(('127.0.0.1', port)) != 0:
                return port
        port += 1
    return start_port

def _read_chrome_devtools_info(port: int, timeout: float = 1.5) -> dict | None:
    """Return DevTools metadata when the port belongs to a debuggable Chrome instance."""
    url = f"http://127.0.0.1:{port}/json/version"
    try:
        # Bypass the app's global proxy settings for localhost DevTools checks.
        opener = urllib_request.build_opener(urllib_request.ProxyHandler({}))
        request = urllib_request.Request(url, headers={"Host": f"127.0.0.1:{port}"})
        with opener.open(request, timeout=timeout) as response:
            payload = response.read().decode("utf-8", errors="replace")
        data = json.loads(payload)
    except (urllib_error.URLError, TimeoutError, OSError, ValueError):
        return None

    browser_name = str(data.get("Browser", "")).lower()
    if not browser_name.startswith("chrome/"):
        return None
    if not data.get("webSocketDebuggerUrl"):
        return None
    return data


def _is_port_listening(port: int, timeout: float = 1.0) -> bool:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.settimeout(timeout)
        return s.connect_ex(("127.0.0.1", port)) == 0


def _find_chrome_executable() -> str | None:
    paths = [
        os.path.expandvars(r"%ProgramFiles%\Google\Chrome\Application\chrome.exe"),
        os.path.expandvars(r"%ProgramFiles(x86)%\Google\Chrome\Application\chrome.exe"),
        os.path.expandvars(r"%LocalAppData%\Google\Chrome\Application\chrome.exe"),
    ]
    return next((p for p in paths if os.path.exists(p)), None)


def _patch_distutils_compat() -> None:
    """Provide a distutils compatibility shim for packages that still import it on Python 3.12+."""
    try:
        import distutils.version  # noqa: F401
        return
    except ModuleNotFoundError:
        pass

    try:
        import setuptools._distutils as setuptools_distutils
        import setuptools._distutils.version as setuptools_distutils_version

        sys.modules.setdefault("distutils", setuptools_distutils)
        sys.modules.setdefault("distutils.version", setuptools_distutils_version)
    except Exception:
        pass


def _iter_site_packages() -> list[Path]:
    """Return likely site-packages locations for the current and adjacent Python installs."""
    candidates: list[Path] = []

    try:
        import site

        for path in site.getsitepackages():
            candidates.append(Path(path))
        user_site = site.getusersitepackages()
        if user_site:
            candidates.append(Path(user_site))
    except Exception:
        pass

    candidates.append(Path(sys.prefix) / "Lib" / "site-packages")
    candidates.append(Path(sys.base_prefix) / "Lib" / "site-packages")

    python_installs_root = Path.home() / "AppData" / "Local" / "Programs" / "Python"
    if python_installs_root.exists():
        candidates.extend(python_installs_root.glob("Python*/Lib/site-packages"))

    seen: set[str] = set()
    existing: list[Path] = []
    for candidate in candidates:
        candidate_str = str(candidate)
        if candidate.exists() and candidate_str not in seen:
            seen.add(candidate_str)
            existing.append(candidate)
    return existing


def _import_undetected_chromedriver():
    """Import undetected_chromedriver with compatibility fallbacks."""
    _patch_distutils_compat()

    try:
        return importlib.import_module("undetected_chromedriver")
    except Exception as first_exc:
        last_exc = first_exc

    for site_packages in _iter_site_packages():
        package_dir = site_packages / "undetected_chromedriver"
        if not package_dir.exists():
            continue

        site_packages_str = str(site_packages)
        if site_packages_str not in sys.path:
            sys.path.append(site_packages_str)

        _patch_distutils_compat()
        try:
            return importlib.import_module("undetected_chromedriver")
        except Exception as exc:
            last_exc = exc

    raise RuntimeError(
        "undetected-chromedriver could not be imported in the active runtime. "
        f"Interpreter: {sys.executable}. Last import error: {last_exc}"
    ) from last_exc


def launch_chrome_debug(port: int = 9222, startup_timeout: float = 15.0) -> tuple[int | None, str | None]:
    """
    Ensure a debuggable Chrome instance is available and return `(port, error_message)`.
    """
    import subprocess

    # 1. Reuse an existing DevTools endpoint when it is healthy.
    if _read_chrome_devtools_info(port):
        logger.info(f"Chrome debug session detected on port {port}")
        return port, None

    # 2. If the preferred port is already occupied, do not launch another Chrome instance.
    if _is_port_listening(port):
        message = (
            f"Port {port} is already occupied, but Chrome DevTools is not reachable there. "
            f"If Chrome is already open, close it and reopen it with "
            f"'--remote-debugging-port={port}' and a separate '--user-data-dir', then try again."
        )
        logger.error(message)
        return None, message

    # 3. Identify Chrome Path
    chrome_path = _find_chrome_executable()
    
    if not chrome_path:
        logger.error("Google Chrome installation not found.")
        return None, "Google Chrome installation not found."

    try:
        # 4. Launch with remote debugging and wait until DevTools is ready.
        profile_dir = DATA_ROOT / f"ChromeAutomator-{port}"
        profile_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Launching outside Chrome in debug mode: {chrome_path} on port {port}")
        subprocess.Popen([
            chrome_path,
            f"--remote-debugging-port={port}",
            "--user-data-dir=" + str(profile_dir),
            "--no-first-run",
            "--no-default-browser-check",
            "about:blank",
        ])

        deadline = time.time() + startup_timeout
        while time.time() < deadline:
            if _read_chrome_devtools_info(port):
                return port, None
            time.sleep(0.5)

        message = f"Chrome started, but DevTools did not become ready on port {port} within {startup_timeout}s."
        logger.error(message)
        return None, message
    except Exception as e:
        logger.error(f"Failed to launch Chrome: {e}")
        return None, f"Failed to launch Chrome: {e}"


def connect_to_chrome_debug(port: int = 9222, connect_timeout: float = 12.0):
    """Attach Selenium to an existing Chrome DevTools session."""
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options

    deadline = time.time() + connect_timeout
    last_error = None

    while time.time() < deadline:
        if not _read_chrome_devtools_info(port):
            time.sleep(0.5)
            continue

        chrome_options = Options()
        chrome_options.add_experimental_option("debuggerAddress", f"127.0.0.1:{port}")
        try:
            return webdriver.Chrome(options=chrome_options)
        except Exception as exc:
            last_error = exc
            time.sleep(0.75)

    if last_error is not None:
        raise last_error
    raise RuntimeError(f"Chrome DevTools was not ready on 127.0.0.1:{port}")


def _normalize_browser_url(url: str) -> tuple[str, str, str]:
    """Return a normalized `(host, path, query)` tuple for fuzzy tab matching."""
    if not url:
        return "", "", ""

    parsed = urlsplit(url.strip())
    host = (parsed.netloc or "").lower()
    path = unquote((parsed.path or "").rstrip("/")).lower()
    query = unquote(parsed.query or "").lower()
    return host, path, query


def _score_page_match(current_url: str, target_url: str) -> int:
    """Score how closely the current browser URL matches the requested URL."""
    # Standardize both URLs to lower/forward-slashes BEFORE splitting
    curr = current_url.lower().replace('\\', '/')
    targ = target_url.lower().replace('\\', '/')
    
    current_host, current_path, current_query = _normalize_browser_url(curr)
    target_host, target_path, target_query = _normalize_browser_url(targ)

    score = 0

    if current_host or target_host:
        if current_host == target_host:
            score += 20
        elif target_host in current_host or current_host in target_host:
            score += 10
        else:
            return 0
    else:
        # Both are local files (no host)
        score += 20

    if current_path and target_path:
        # Standardize for comparison - lowercase and forward slashes for cross-platform and drive letter consistency
        cp = current_path.replace('\\', '/').rstrip('/').lower()
        tp = target_path.replace('\\', '/').rstrip('/').lower()
        
        if cp == tp:
            score += 50
        elif cp.endswith(tp) or tp.endswith(cp):
            score += 35
        elif tp in cp or cp in tp:
            score += 20

    if target_query:
        if current_query == target_query:
            score += 20
        elif current_query and (target_query in current_query or current_query in target_query):
            score += 10

    if current_url.strip().lower() == target_url.strip().lower():
        score += 20

    return score


def _extract_url_tokens(url: str) -> set[str]:
    """Extract stable URL tokens so matching survives query/hash/session noise."""
    host, path, query = _normalize_browser_url(url)
    # Special handling for local files to make matching more robust (case-insensitive and slash-agnostic)
    if not host and path.strip():
         raw = path.lower().replace('\\', '/').strip('/')
    else:
         raw = " ".join(part for part in (host, path, query) if part).lower()
    
    tokens = {token for token in re.split(r"[^a-z0-9]+", raw) if len(token) >= 3}
    return tokens


def safe_get_tab_ids(driver) -> list:
    """Version-agnostic and object-agnostic way to get tab IDs for DrissionPage."""
    try:
        if hasattr(driver, 'tab_ids'): return list(driver.tab_ids)
        if hasattr(driver, 'tabs'): return list(driver.tabs)
        # If driver is a Tab, try getting from browser
        if hasattr(driver, 'browser') and hasattr(driver.browser, 'tab_ids'):
            return list(driver.browser.tab_ids)
        if hasattr(driver, 'page') and hasattr(driver.page, 'tab_ids'):
            return list(driver.page.tab_ids)
    except:
        pass
    return []

def switch_to_matching_page(driver, target_url, timeout=30):
    """
    Tries to find an open tab that matches target_url.
    Returns (success, current_url_or_error)
    """
    target_url = (target_url or "").strip()
    if not target_url:
        if hasattr(driver, 'url'):
            return True, (driver.url or "").strip()
        return True, (driver.current_url or "").strip()

    # Check if driver is DrissionPage (ChromiumPage)
    is_drission = hasattr(driver, 'tab_ids') or "DrissionPage" in str(type(driver))
    
    def _std(u):
        if not u: return ""
        u = u.replace('\\', '/').lower().strip()
        if 'file:/' in u:
             u = u.split('file:/')[-1].lstrip('/')
        return u.rstrip('/')

    target_std = _std(target_url)
    target_tokens = _extract_url_tokens(target_url)
    
    deadline = time.time() + timeout
    best_handle = None
    best_score = 0
    best_url = ""

    while time.time() < deadline:
        if is_drission:
            try:
                for tab_id in safe_get_tab_ids(driver):
                    try:
                        tab = driver.get_tab(tab_id)
                        curr_url = (tab.url or "").strip()
                        curr_std = _std(curr_url)
                        
                        score = _score_page_match(curr_url, target_url)
                        if target_std and curr_std and (target_std == curr_std or target_std in curr_std or curr_std in target_std):
                            score = max(score, 70)

                        if target_tokens:
                            overlap = target_tokens & _extract_url_tokens(curr_url)
                            score += min(len(overlap) * 8, 40)
                        
                        if score > best_score:
                            best_score, best_handle, best_url = score, tab_id, curr_url
                        
                        if score >= 60:
                            # For DrissionPage, the tab object is the one to return as 'driver'
                            return True, tab
                    except Exception:
                        continue
            except Exception as e:
                if "10061" in str(e):
                    time.sleep(1)
                    continue
        else:
            for handle in list(driver.window_handles):
                try:
                    driver.switch_to.window(handle)
                    curr_url = (driver.current_url or "").strip()
                    curr_std = _std(curr_url)
                    
                    score = _score_page_match(curr_url, target_url)
                    if target_std and curr_std and (target_std == curr_std or target_std in curr_std or curr_std in target_std):
                        score = max(score, 70)

                    if target_tokens:
                        overlap = target_tokens & _extract_url_tokens(curr_url)
                        score += min(len(overlap) * 8, 40)
                    
                    if score > best_score:
                        best_score, best_handle, best_url = score, handle, curr_url
                    
                    if score >= 60:
                        return True, curr_url
                except Exception:
                    continue
        
        time.sleep(0.5)

    if best_handle and best_score > 0:
        try:
            if is_drission:
                return True, driver.get_tab(best_handle)
            else:
                driver.switch_to.window(best_handle)
                return True, (driver.current_url or best_url).strip()
        except Exception:
            pass

    # Fallback diagnostics
    visible_urls = []
    if is_drission:
        try:
            for tab_id in safe_get_tab_ids(driver):
                try: visible_urls.append(driver.get_tab(tab_id).url or "")
                except Exception: continue
        except Exception: pass
    else:
        for handle in list(driver.window_handles):
            try:
                driver.switch_to.window(handle)
                visible_urls.append(driver.current_url or "")
            except Exception: continue

    url_summary = "; ".join([u for u in visible_urls if u][:5]) or "no readable tabs"
    return False, f"Could not find an open tab matching {target_url}. Open tabs seen: {url_summary}"


def launch_undetected_chrome(
    start_url: str = "",
    profile_name: str = "ChromeAutomatorUC",
    download_dir: Path | None = None,
):
    """Launch an undetected-chromedriver session with a persistent Chrome profile."""
    uc = _import_undetected_chromedriver()

    chrome_path = _find_chrome_executable()
    if not chrome_path:
        raise RuntimeError("Google Chrome installation not found.")

    profile_dir = DATA_ROOT / profile_name
    profile_dir.mkdir(parents=True, exist_ok=True)

    options = uc.ChromeOptions()
    options.add_argument(f"--user-data-dir={profile_dir}")
    options.add_argument("--no-first-run")
    options.add_argument("--no-default-browser-check")
    options.add_argument("--disable-blink-features=AutomationControlled")

    if download_dir:
        download_dir.mkdir(parents=True, exist_ok=True)
        prefs = {
            "download.default_directory": str(download_dir),
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "safebrowsing.enabled": True,
        }
        options.add_experimental_option("prefs", prefs)

    logger.info(f"Launching undetected Chrome session with profile: {profile_dir}")
    driver = uc.Chrome(
        options=options,
        browser_executable_path=chrome_path,
        user_data_dir=str(profile_dir),
        use_subprocess=True,
    )

    if start_url:
        driver.get(start_url)

    return driver


def launch_managed_selenium_chrome(
    start_url: str = "",
    profile_name: str = "ChromeAutomatorManaged",
    download_dir: Path | None = None,
):
    """Launch a normal Selenium-managed Chrome session with a persistent profile (no debug port)."""
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options

    chrome_path = _find_chrome_executable()
    if not chrome_path:
        raise RuntimeError("Google Chrome installation not found.")

    profile_dir = DATA_ROOT / profile_name
    profile_dir.mkdir(parents=True, exist_ok=True)

    options = Options()
    options.binary_location = chrome_path
    options.add_argument(f"--user-data-dir={profile_dir}")
    options.add_argument("--no-first-run")
    options.add_argument("--no-default-browser-check")
    options.add_argument("--disable-blink-features=AutomationControlled")

    if download_dir:
        download_dir.mkdir(parents=True, exist_ok=True)
        prefs = {
            "download.default_directory": str(download_dir),
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "safebrowsing.enabled": True,
        }
        options.add_experimental_option("prefs", prefs)

    logger.info(f"Launching managed Selenium Chrome session with profile: {profile_dir}")
    driver = webdriver.Chrome(options=options)
    if start_url:
        driver.get(start_url)
    return driver


def launch_drission_chrome(
    start_url: str = "",
    port: int = 9222,
    profile_name: str | None = "ChromeAutomatorDirect",
    download_dir: Path | None = None,
):
    """Launch a DrissionPage Chromium session for stealthy automation."""
    try:
        from DrissionPage import ChromiumPage, ChromiumOptions
    except ImportError:
        raise RuntimeError("DrissionPage is not installed. Please install it with 'pip install DrissionPage'.")

    co = ChromiumOptions()
    chrome_path = _find_chrome_executable()
    if chrome_path:
        co.set_browser_path(chrome_path)
    
    if profile_name:
        profile_dir = DATA_ROOT / profile_name
        profile_dir.mkdir(parents=True, exist_ok=True)
        co.set_user_data_path(str(profile_dir))
        logger.info(f"Launching DrissionPage Chrome session with profile: {profile_dir}")
    else:
        logger.info(f"Attaching DrissionPage to existing Chrome on port {port}")

    co.set_local_port(port)
    
    # DrissionPage specific stealth settings
    co.set_argument("--no-first-run")
    co.set_argument("--no-default-browser-check")
    
    if download_dir:
        download_dir.mkdir(parents=True, exist_ok=True)
        co.set_download_path(str(download_dir))

    page = ChromiumPage(co)
    
    if start_url:
        page.get(start_url)
        
    return page


def get_automation_driver(
    start_url: str = "",
    port: int = 9222,
    profile_name: str = "ChromeAutomatorUC",
    download_dir: Path | None = None,
    allow_browser_launch: bool = True,
    use_direct_mode: bool = False,
    allow_debug: bool = True,
):
    """
    Reuse an existing debug Chrome session when available.
    Otherwise, prefer DrissionPage (Direct Mode) or undetected-chromedriver (Debug/Stealth Mode).
    """
    # 1. If Direct Mode is explicitly requested
    if use_direct_mode:
        # We NO LONGER launch a new browser. We only attach to an existing one on 9222.
        logger.info("Direct Mode: Attaching to existing Chrome on port 9222...")
        try:
            # profile_name=None ensures we attach rather than launch a new profile
            driver = launch_drission_chrome(
                start_url=start_url,
                port=9222,
                profile_name=None,
                download_dir=download_dir
            )
            if driver:
                logger.info("Successfully attached to existing Chrome session.")
                return driver, "direct"
        except Exception as e:
            logger.warning(f"Could not attach to 9222: {e}")
        
        # If we reach here, connection failed. We do not fallback.
        raise RuntimeError(
            "Direct Mode Failed: No open Chrome window found on port 9222.\n"
            "Solution: Close all Chrome windows and restart Chrome from your terminal with:\n"
            "chrome.exe --remote-debugging-port=9222"
        )

    # 2. Reuse an existing debug Chrome session when available.
    if allow_debug and _read_chrome_devtools_info(port):
        logger.info(f"Reusing existing Chrome debug session on port {port}")
        return connect_to_chrome_debug(port=port), "debug"

    if not allow_browser_launch:
        raise RuntimeError("Browser launch is disabled and Direct/Debug attach is not available.")

    uc_error = None
    # Prefer launching a managed stealth session when we are allowed to launch,
    # even if start_url is empty (the caller may navigate manually).
    try:
        return launch_undetected_chromedriver(
            start_url=start_url,
            profile_name=profile_name,
            download_dir=download_dir,
        ), "undetected"
    except Exception as exc:
        uc_error = str(exc)
        logger.warning(f"Undetected Chrome launch failed: {exc}")

    # Fallback: launch a standard managed Selenium Chrome (no debug port).
    managed_error = None
    try:
        return launch_managed_selenium_chrome(
            start_url=start_url,
            profile_name=profile_name,
            download_dir=download_dir,
        ), "managed"
    except Exception as exc:
        managed_error = str(exc)
        logger.warning(f"Managed Selenium Chrome launch failed: {exc}")

    debug_error = None
    if allow_debug:
        debug_port, debug_error = launch_chrome_debug(port=port)
        if debug_port:
            return connect_to_chrome_debug(port=debug_port), "debug"

    error_parts = []
    if uc_error:
        error_parts.append(f"Undetected-Chromedriver failed: {uc_error}")
    if managed_error:
        error_parts.append(f"Managed Selenium Chrome failed: {managed_error}")
    if debug_error:
        error_parts.append(debug_error)
    if not start_url:
        error_parts.append(
            "Could not launch a managed browser profile and no reusable debug Chrome session was found. "
            "If your environment blocks managed launch, use Direct Mode with --remote-debugging-port=9222 "
            "or provide the target GeM URL."
        )

    raise RuntimeError(" ".join(error_parts).strip())

# ─── Other Utilities ───────────────────────────────────────────────────────────
def list_visible_elements(xpath, driver, driver_mode, root=None):
    """Find visible elements across either Selenium or DrissionPage."""
    if driver_mode == "direct":
        scope = root if root is not None else driver
        try:
            # DrissionPage uses 'xpath:' prefix for explicit xpath
            els = scope.eles(f'xpath:{xpath}')
            return [el for el in els if el]
        except Exception:
            return []

    from selenium.webdriver.common.by import By
    scope = root if root is not None else driver
    try:
        return [el for el in scope.find_elements(By.XPATH, xpath) if el.is_displayed()]
    except Exception:
        return []

def get_url(driver, driver_mode):
    """Get current URL across drivers."""
    if driver_mode == "direct":
        return (driver.url or "").strip()
    return (driver.current_url or "").strip()

def run_script(driver, driver_mode, script, *args):
    """Execute javascript across drivers."""
    if driver_mode == "direct":
        return driver.run_js(script, *args)
    return driver.execute_script(script, *args)

def is_same_window(driver, driver_mode, handle):
    """Check window handle match across drivers."""
    if driver_mode == "direct":
        return driver.tab_id == handle
    return driver.current_window_handle == handle

def _move_mouse_stealthy(driver, target_el):
    """Move mouse along a Bezier-like curve with jitter to mimic human physiology."""
    try:
        from random import randint, uniform
        # Get start/end coords
        # DrissionPage .rect() returns coordinates
        rect = target_el.rect
        end_x = rect.x + (rect.width / 2) + uniform(-5, 5)
        end_y = rect.y + (rect.height / 2) + uniform(-5, 5)
        
        # Simple 2-point interpolation with jitter
        # In a full Bezier implementation we'd calculate intermediate points
        # For now, we'll just move to the target with a slight delay and hover
        driver.actions.move_to(target_el)
        time.sleep(uniform(0.1, 0.3))
    except Exception:
        pass

def safe_click(driver, driver_mode, el):
    """Robust click helper across drivers with stealth emulation."""
    from random import uniform
    if driver_mode == "direct":
        try:
            el.scroll.to_see()
            _move_mouse_stealthy(driver, el)
            time.sleep(uniform(0.1, 0.4))
            el.click()
            return
        except Exception:
            pass
        try:
            driver.run_js("arguments[0].click();", el)
            return
        except Exception:
            pass
        return

    from selenium.webdriver.common.action_chains import ActionChains
    try:
        run_script(driver, driver_mode, "arguments[0].scrollIntoView({block: 'center', inline: 'center'});", el)
        time.sleep(0.3)
        # Try pointer interaction first
        ActionChains(driver).move_to_element(el).click().perform()
        return
    except Exception:
        pass
    try:
        el.click()
        return
    except Exception:
        pass
    run_script(driver, driver_mode, "arguments[0].click();", el)

def set_value(driver, driver_mode, el, value):
    """Set input value with human-like typing delays (50ms-200ms per char)."""
    from random import uniform
    if driver_mode == "direct":
        try:
            el.clear()
            for char in str(value):
                el.input(char)
                time.sleep(uniform(0.05, 0.2))
            return
        except Exception:
            pass
        try:
            driver.run_js("arguments[0].value = arguments[1];", el, value)
            return
        except Exception:
            pass
        return

    # Selenium mode
    try:
        el.clear()
        for char in str(value):
            el.send_keys(char)
            time.sleep(uniform(0.05, 0.15))
    except Exception:
        driver.execute_script("arguments[0].value = arguments[1];", el, value)

def get_frame(driver, driver_mode, locator):
    """Switch to or get an iframe across drivers."""
    if driver_mode == "direct":
        try:
            # DrissionPage .get_frame() returns the frame object
            return driver.get_frame(locator)
        except Exception:
            return None
    
    # For Selenium...
    from selenium.webdriver.common.by import By
    try:
        frame_el = driver.find_element(By.XPATH, locator) if "/" in str(locator) else driver.find_element(By.ID, locator)
        driver.switch_to.frame(frame_el)
        return driver
    except Exception:
        return None

def run_automation_steps(driver, driver_mode, steps, emit_callback=None):
    """
    Executes a list of automation steps in a data-driven way, mimicking Automa.
    Step Schema: {'type': 'click'|'type'|'wait'|'scroll', 'selector': '...', 'value': '...', 'delay': 0.5}
    """
    import time
    from random import uniform
    for step in steps:
        stype = step.get("type")
        sel = step.get("selector")
        val = step.get("value")
        delay = step.get("delay", uniform(0.5, 1.5))
        
        try:
            if stype == "wait":
                time.sleep(delay)
                continue
                
            el = None
            if sel:
                els = list_visible_elements(sel, driver, driver_mode)
                if not els:
                    if step.get("mandatory", True):
                        raise Exception(f"Step {stype} failed: Element not found '{sel}'")
                    continue
                el = els[0]
            
            if stype == "click":
                safe_click(driver, driver_mode, el)
            elif stype == "type":
                set_value(driver, driver_mode, el, val)
            elif stype == "scroll":
                if driver_mode == "direct":
                    el.scroll.to_see()
                else:
                    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", el)
            
            time.sleep(delay)
        except Exception as e:
            if emit_callback: 
                emit_callback("info", {"message": f"Step Error: {e}"})
            if step.get("mandatory", True): raise e
    return True
